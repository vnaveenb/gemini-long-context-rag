"""DOCX text extraction using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401

from src.logger import get_logger
from src.models.document import (
    DocumentFormat,
    DocumentMetadata,
    ExtractedContent,
    PageContent,
    Section,
)

logger = get_logger(__name__)

_HEADING_STYLES = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Title": 1,
}


def extract_docx(file_path: str | Path) -> ExtractedContent:
    """Extract text, sections, and metadata from a DOCX file."""
    file_path = Path(file_path)
    logger.info("Extracting DOCX", path=str(file_path))

    errors: list[str] = []

    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        logger.error("Failed to open DOCX", path=str(file_path), error=str(exc))
        return ExtractedContent(
            filename=file_path.name,
            format=DocumentFormat.DOCX,
            extraction_errors=[f"Failed to open DOCX: {exc}"],
        )

    sections: list[Section] = []
    all_text_parts: list[str] = []
    current_section_text: list[str] = []
    page_number = 1  # DOCX doesn't have native page numbers; estimate

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            all_text_parts.append("")
            current_section_text.append("")
            continue

        style_name = para.style.name if para.style else ""
        level = _HEADING_STYLES.get(style_name)

        if level is not None:
            # Close previous section content
            if sections:
                sections[-1].content = "\n".join(current_section_text).strip()
                current_section_text = []

            sections.append(
                Section(
                    title=text,
                    level=level,
                    page_start=page_number,
                )
            )
        else:
            current_section_text.append(text)

        all_text_parts.append(text)

        # Rough page estimation (~3000 chars per page)
        if sum(len(t) for t in all_text_parts[-20:]) > 3000:
            page_number += 1

    # Close last section
    if sections:
        sections[-1].content = "\n".join(current_section_text).strip()
        for i, sec in enumerate(sections):
            sec.page_end = sections[i + 1].page_start if i + 1 < len(sections) else page_number

    # Extract tables as additional text
    for i, table in enumerate(doc.tables):
        try:
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            table_text = f"\n[Table {i + 1}]\n" + "\n".join(rows_text)
            all_text_parts.append(table_text)
        except Exception as exc:
            errors.append(f"Table {i + 1}: {exc}")

    raw_text = "\n".join(all_text_parts)

    # Create pseudo-pages
    words = raw_text.split()
    page_size = 500  # words per page
    pages = [
        PageContent(page_number=idx + 1, text=" ".join(words[i : i + page_size]))
        for idx, i in enumerate(range(0, len(words), page_size))
    ]

    core = doc.core_properties
    metadata = DocumentMetadata(
        author=core.author,
        title=core.title,
        subject=core.subject,
        creation_date=core.created,
        modification_date=core.modified,
        page_count=len(pages),
        word_count=len(words),
    )

    content = ExtractedContent(
        filename=file_path.name,
        format=DocumentFormat.DOCX,
        metadata=metadata,
        pages=pages,
        sections=sections,
        raw_text=raw_text,
        extraction_errors=errors,
    )
    logger.info(
        "DOCX extraction complete",
        pages=len(pages),
        sections=len(sections),
        words=metadata.word_count,
    )
    return content
